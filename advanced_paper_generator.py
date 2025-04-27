"""
高级论文生成器 - 专注于生成长篇幅、高质量且难以被AI检测的学术论文
"""

import os
import time
import json,re
import random
import requests
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert  # 添加新的导入
#os.environ['OPENAI_API_KEY'] = 'sk-QXFhTtcE8Eo9NpA47501AdC0Da3840Ee96DfDc693400B9Ec'
os.environ['OPENAI_API_KEY'] ='sk-5vhOBlyfX42oyjHo9d5929F5A46342C4A0C3D49446459a3c'
os.environ['OPENAI_API_URL'] = 'https://apione.zen-x.com.cn/api/v1'
os.environ['OPENAI_API_MODEL'] = 'gpt-4o'
os.environ['ALI_API_KEY'] = 'sk-ed55536fef274e5ca8699c0764b3b624'
os.environ['ALI_API_URL'] = 'https://dashscope.aliyuncs.com/compatible-mode/v1'
os.environ['ALI_API_MODEL'] = 'qwen-max-latest'
os.environ['HUOSHAN_API_KEY'] = 'd8d0c6d8-d566-4060-8fbb-c6f517383b12'
os.environ['HUOSAHN_API_URL'] = 'https://ark.cn-beijing.volces.com/api/v3'
os.environ['HUOSHAN_API_MODEL'] = 'Doubao-1.5-pro-32k'
# API配置
API_KEY = os.environ.get("OPENAI_API_KEY")
API_URL = os.environ.get("OPENAI_API_URL")
API_MODEL = os.environ.get("OPENAI_API_MODEL")  # 使用GPT-4.1或等效模型

class AdvancedPaperGenerator:
    """高级学术论文生成器类"""
    
    # 定义可用的模型配置
    MODEL_CONFIGS = {
        'gpt4o': {
            'api_key': os.environ.get('OPENAI_API_KEY'),
            'api_url': os.environ.get('OPENAI_API_URL'),
            'api_model': "gpt-4o"
        },
        'qwen': {
            'api_key': os.environ.get('ALI_API_KEY'),
            'api_url': os.environ.get('ALI_API_URL'),
            'api_model': 'qwen-max-latest'
        },
        'doubao': {
            'api_key': os.environ.get('HUOSHAN_API_KEY'),
            'api_url': os.environ.get('HUOSAHN_API_URL'),
            'api_model': 'ep-20250327154845-jjhgt'
        },
        'deepseek': {
            'api_key': os.environ.get('HUOSHAN_API_KEY'),
            'api_url': os.environ.get('HUOSAHN_API_URL'),
            'api_model': "ep-20250327155938-zngrc"
        }
    }
    
    def __init__(self, model_name: str = 'gpt4o'):
        """
        初始化生成器
        :param model_name: 使用的模型名称，可选值：'gpt4o', 'qwen', 'doubao', 'deepseek'
        """
        if model_name not in self.MODEL_CONFIGS:
            raise ValueError(f"不支持的模型名称: {model_name}。可用模型: {list(self.MODEL_CONFIGS.keys())}")
            
        config = self.MODEL_CONFIGS[model_name]
        self.api_key = config['api_key']
        self.api_url = config['api_url']
        self.api_model = config['api_model']
        self.current_model = model_name
        
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        self.topic = None
        self.wait_range = (1, 3)  # API调用间隔时间范围(秒)
    def _clean_content(self, content: str) -> str:
        """清理API返回的内容"""
        # 移除可能的markdown标记
        if content.startswith('```json'):
            content = content[7:]
        if content.endswith('```'):
            content = content[:-3]
        content = content.strip()
        
        # 处理可能的JSON字符串转义问题
        content = content.replace('\\"', '"')
        
        # 移除内容中的JSON键名
        content = content.replace('"title":', '')
        content = content.replace('"abstract":', '')
        content = content.replace('"content":', '')
        content = content.replace('"reference":', '')
        content = content.replace('"references":', '')
        
        # 移除多余的引号和大括号
        content = content.strip('{}')
        
        return content
    def _call_api(self, messages: List[Dict], temperature: float = 0.5, 
                 max_tokens: int = 8192,response_format: str = "text") -> Optional[str]:
        """API调用方法"""
        data = {
            "model": self.api_model,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens,
            "response_format": { "type": response_format }
        }
        try:
            time.sleep(random.uniform(*self.wait_range))
            response = requests.post(
                f"{self.api_url}/chat/completions",
                headers=self.headers,
                data=json.dumps(data),
                timeout=120
            )
            
            if response.status_code == 200:
                print(response.json()["choices"][0]["message"]["content"].strip())
                return response.json()["choices"][0]["message"]["content"].strip()
            else:
                print(f"API调用失败: HTTP {response.status_code}")
                print(f"错误信息: {response.text}")
                exit(1)
                
        except Exception as e:
            print(f"API调用异常: {str(e)}")
            exit(1)

    def generate_references(self, topic: str, count: int = 25) -> List[Dict]:
        """生成IEEE格式的参考文献"""
        try:
            # 首先获取相关的研究方向和关键词
            direction_prompt = f"""For the topic "{topic}", list 5 specific research directions and their related keywords.
            Return as JSON array:
            [
                {{
                    "direction": "specific research direction",
                    "keywords": ["keyword1", "keyword2", "keyword3"]
                }}
            ]"""
            
            messages = [
                {
                    "role": "system",
                    "content": "You are a research topic analyzer helping to identify specific research directions and keywords."
                },
                {"role": "user", "content": direction_prompt}
            ]
            
            # 获取研究方向
            directions_response = self._call_api(messages, temperature=0.3, response_format="json_object")
            directions_data = json.loads(directions_response)
            
            # 确保我们有有效的方向列表
            if isinstance(directions_data, dict):
                directions = directions_data.get("directions", [])
            else:
                directions = directions_data
                
            if not directions:  # 如果方向列表为空
                print("未能获取有效的研究方向")
                return []
                
            # 基于研究方向生成参考文献
            references = []
            # 确保每个方向至少生成一个引用，同时平均分配总数
            refs_per_direction = max(1, count // len(directions))
            extra_refs = count % len(directions)  # 处理不能整除的情况
            
            for i, direction in enumerate(directions):
                # 为最后几个方向多分配一个引用，以达到总数
                current_refs_count = refs_per_direction + (1 if i < extra_refs else 0)
                
                ref_prompt = f"""Create {current_refs_count} example academic reference entries about {direction['direction']}.
                Include keywords like: {', '.join(direction['keywords'])}
                
                Format as JSON array with example references:
                [
                    {{
                        "ref_number": "number",
                        "authors": ["A. LastName1", "B. LastName2"],
                        "year": "between 2015-2025",
                        "title": "title using some keywords",
                        "venue": "relevant journal or conference name",
                        "volume": "vol number",
                        "issue": "issue number",
                        "pages": "page range",
                        "doi": "optional doi"
                    }}
                ]"""
                
                messages = [
                    {
                        "role": "system",
                        "content": "You are helping to create example reference formats for educational purposes."
                    },
                    {"role": "user", "content": ref_prompt}
                ]
                
                try:
                    ref_response = self._call_api(messages, temperature=0.4, response_format="json_object")
                    ref_data = json.loads(ref_response)
                    
                    if isinstance(ref_data, dict):
                        direction_refs = ref_data.get("references", [])
                    else:
                        direction_refs = ref_data
                        
                    # 更新引用编号
                    start_num = len(references) + 1
                    for i, ref in enumerate(direction_refs, start=start_num):
                        ref["ref_number"] = i
                        references.append(ref)
                        
                    if len(references) >= count:
                        break
                        
                except Exception as e:
                    print(f"生成单个方向的参考文献时出错: {str(e)}")
                    continue
            
            # 验证并返回结果
            if not references:
                print("未能生成任何有效的参考文献")
                return []
                
            # 确保不超过请求的数量
            references = references[:count]
            
            # 验证格式
            valid_references = []
            for ref in references:
                required_fields = ["authors", "year", "title", "venue"]
                if all(field in ref for field in required_fields):
                    valid_references.append(ref)
            
            return valid_references
            
        except Exception as e:
            print(f"生成参考文献时出错: {str(e)}")
            return []



    def format_references(self, references: List[Dict]) -> str:
        """格式化IEEE样式的参考文献"""
        formatted_refs = []
        for ref in references:
            try:
                # 作者格式化
                authors = ref.get("authors", [])
                if len(authors) > 3:
                    authors_str = ", ".join(authors[:3]) + " et al."
                else:
                    authors_str = ", ".join(authors[:-1])
                    if authors:
                        authors_str += ", and " + authors[-1]
            
                # 基本引用信息
                ref_str = f"[{ref.get('ref_number', '')}] {authors_str}, "
                ref_str += f"\"{ref.get('title', '')},\" "
                ref_str += f"{ref.get('venue', '')}"
                
                # 可选信息
                if ref.get('volume'):
                    ref_str += f", vol. {ref.get('volume', '')}"
                if ref.get('issue'):
                    ref_str += f", no. {ref.get('issue', '')}"
                if ref.get('pages'):
                    ref_str += f", pp. {ref.get('pages', '')}"
                ref_str += f", {ref.get('year', '')}"
                if ref.get('doi'):
                    ref_str += f", doi: {ref.get('doi', '')}"
                
                formatted_refs.append(ref_str + ".")
            
            except Exception as e:
                print(f"格式化参考文献时出错: {str(e)}")
                continue
        
        return "\n".join(formatted_refs)

    def save_paper_to_docx(self, output_file: str) -> None:
        """保存为docx文件"""
        try:
            paper = self.generate_full_paper(self.topic)
            if not paper:
                print("错误：没有内容可保存")
                exit(1)
            
            doc = Document()
            
            # 设置文档样式
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # 标题
            title = doc.add_heading(level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run(paper.get("title", ""))
            title_run.font.size = Pt(16)
            
            # 添加内容部分
            for section in paper.get("sections", []):
                # 添加带有标题样式的章节标题
                heading = doc.add_heading(section.get("name", ""), level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 处理内容，移除可能的重复标题
                content = section.get("content", "")
                
                # 移除各种可能的标题格式
                patterns_to_remove = [
                    f"**{section.get('name', '')}**",  # Markdown格式
                    f"# {section.get('name', '')}",    # Markdown格式
                    f"## {section.get('name', '')}",    # Markdown格式
                    f"### {section.get('name', '')}",    # Markdown格式
                    section.get("name", ""),           # 普通文本格式
                    f"{section.get('name', '')}:",     # 带冒号的格式
                ]
                
                # 清理内容开头的标题
                content_lines = content.split('\n')
                while content_lines and any(content_lines[0].strip().lower() == pattern.lower() 
                                          for pattern in patterns_to_remove):
                    content_lines.pop(0)
                
                # 重新组合内容并添加到文档
                cleaned_content = '\n'.join(content_lines).strip()
                if cleaned_content:
                    doc.add_paragraph(cleaned_content)
            
            # 添加参考文献
            doc.add_heading("References", level=2)
            refs = self.format_references(paper.get("references", []))
            for ref in refs.split("\n"):
                if ref.strip():
                    p = doc.add_paragraph()
                    p.add_run(ref.strip())
                    p.paragraph_format.first_line_indent = Inches(0.5)
                    p.paragraph_format.left_indent = Inches(0)
            
            doc.save(output_file)
            print(f"\n论文已保存到: {output_file}")
            
        except Exception as e:
            print(f"保存文件时出错: {str(e)}")
            exit(1)

    def save_paper_to_pdf(self, output_file: str) -> None:
        """保存为PDF文件"""
        try:
            # 首先保存为临时docx文件
            temp_docx = output_file.replace('.pdf', '.docx')
            self.save_paper_to_docx(temp_docx)
            
            # 转换为PDF
            try:
                convert(temp_docx)
                
                # 验证PDF文件是否成功创建
                if not os.path.exists(output_file):
                    raise Exception("PDF文件未能成功创建")
                
                print(f"\n论文已保存到: {output_file}")
                
                # 删除临时docx文件
                try:
                    os.remove(temp_docx)
                except Exception as e:
                    print(f"警告：无法删除临时文件 {temp_docx}: {str(e)}")
                           
            except Exception as e:
                print(f"转换PDF时出错: {str(e)}")
                print(f"保留DOCX格式文件作为备份: {temp_docx}")
                return
            
        except Exception as e:
            print(f"保存文件时出错: {str(e)}")
            raise  # 抛出异常而不是exit(1)，让调用者决定如何处理
    def _clean_inline_references(self, content: str) -> str:
        """清理段落中的参考文献"""
        # 移除段落末尾的参考文献列表
        patterns = [
            # 匹配numbered格式的参考文献 [1], [2], etc.
            r'\s*\[\d+[\s*,\s*\d+]*\]\s*\.\s*$',
            # 匹配完整的参考文献条目
            r'\s*\[\d+\]\s+[^[]+?\.\s*$',
            # 匹配参考文献列表
            r'\s*References:?\s*\n.*$',
            # 匹配以数字列表开头的参考文献
            r'\s*\d+\.\s+[A-Z][^.]+\.\s*$',
            # 匹配以作者年份格式的参考文献
            r'\s*\([A-Za-z]+\s+et\s+al\.,\s*\d{4}\)\.\s*$',
            # 匹配多行参考文献
            r'\s*\n\s*References:?\s*\n[\s\S]*$'
        ]
        
        cleaned_content = content
        for pattern in patterns:
            cleaned_content = re.sub(pattern, '', cleaned_content, flags=re.MULTILINE | re.IGNORECASE)
        
        # 移除段落中的引用标记但保留句点
        cleaned_content = re.sub(r'\s*\[\d+[\s*,\s*\d+]*\](?=\.)', '', cleaned_content)
    
        return cleaned_content.strip()
    def generate_section_content(self, section_name: str, topic: str, min_words: int = 2000) -> str:
        """生成单个章节的详细内容"""
        section_prompts = {
            "Abstract": """Write a concise abstract (250 words) that summarizes:
            - The research problem and motivation
            - Key methodology
            - Main results
            - Principal conclusions
            Format: Start each paragraph with two spaces for indentation.
            Do not include references or citations.""",
            
            "Introduction": """Write an introduction that covers:
            - Problem statement and motivation
            - Research objectives
            - Main contributions
            Keep it focused and around 600 words.
            Do not include a reference list at the end.""",
            
            "Related Work": """Review relevant literature covering:
            - Key related research
            - Current limitations
            - How your work differs
            Aim for about 750 words.
            Do not add a reference list at the end.""",
            
            "Methodology": """Describe the research methodology including:
            - Approach and methods used
            - Implementation details
            - Experimental setup
            Target around 1200 words.
            Do not include references at the end.""",
            
            "Results": """Present the main findings:
            - Key results and observations
            - Data analysis
            - Performance metrics
            About 850 words.
            Do not include a reference section.""",
            
            "Discussion": """Discuss the implications:
            - Interpretation of results
            - Limitations
            - Future directions
            Around 500 words.
            Do not add references at the end.""",
            
            "Conclusion": """Conclude with:
            - Summary of contributions
            - Key takeaways
            - Future work
            Brief, about 250 words.
            Do not include references."""
        }

        prompt = f"""Generate the {section_name} section for an academic paper about "{topic}".
        
        {section_prompts.get(section_name, "Generate detailed content.")}
        
        Requirements:
        1. Use academic language
        2. Be specific and technical
        3. Include relevant citations
        4. Stay focused and concise"""

        messages = [
        {
            "role": "system",
            "content": "You are an expert academic writer. Generate clear, focused content without JSON formatting."
        },
        {"role": "user", "content": prompt}
        ]
        
        response = self._call_api(messages, temperature=0.5, max_tokens=min_words,response_format="text")
        return response

    def generate_full_paper(self, topic: str) -> Dict:
        """生成完整论文"""
        try:
            # 生成论文标题
            title_prompt = f"""Generate a concise academic title for a research paper about {topic}.
            The title should be specific and technical, typically 10-15 words."""
            
            title_response = self._call_api(messages=[
                {"role": "system", "content": "Generate a concise academic paper title."},
                {"role": "user", "content": title_prompt}
            ], temperature=0.5, max_tokens=100,response_format="text")

            paper = {
                "title": title_response.strip(),
                "sections": []
            }

            # 定义每个部分的合理字数要求
            section_requirements = [
                ("Abstract", 250),           # 通常250-300字
                ("Introduction", 800),       # 约2页
                ("Related Work", 1050),       # 约3页
                ("Methodology", 1500),       # 约4页
                ("Results", 1100),            # 约3页
                ("Discussion", 800),         # 约2页
                ("Conclusion", 550)          # 约1页
            ]

            # 逐个生成每个部分
            for section_name, min_words in section_requirements:
                print(f"\n生成 {section_name} 部分...")
                content = self.generate_section_content(section_name, topic, min_words)
                
                paper["sections"].append({
                    "name": section_name,
                    "content": content
                })

            # 生成参考文献（通常15-25个引用）
            print("\n生成参考文献...")
            references = self.generate_references(topic, count=20)
            paper["references"] = references

            return paper

        except Exception as e:
            print(f"生成论文时出错: {str(e)}")
            exit(1)

    def generate_single_topic(self) -> Dict:
        """生成单个研究主题"""
        fields = {
            "Computer Science": [
                "Artificial Intelligence", "Machine Learning", "Computer Vision",
                "Natural Language Processing", "Cybersecurity", "Cloud Computing",
                "Internet of Things", "Blockchain", "Software Engineering",
                "Quantum Computing", "Edge Computing", "Big Data Analytics",
                "Human-Computer Interaction", "Distributed Systems", "Network Security",
                "Mobile Computing", "Augmented Reality", "Virtual Reality",
                "Robotics Software", "Database Systems", "Algorithm Design",
                "Computer Graphics", "Parallel Computing", "Web Technologies"
            ],
            "Physics": [
                "Quantum Mechanics", "Particle Physics", "Astrophysics",
                "Condensed Matter", "Optics", "Nuclear Physics",
                "Plasma Physics", "String Theory", "Dark Matter Research",
                "Quantum Field Theory", "Cosmology", "Atomic Physics",
                "Solid State Physics", "Biophysics", "Thermodynamics",
                "Quantum Optics", "Gravitational Physics", "High Energy Physics",
                "Materials Physics", "Nanophysics", "Theoretical Physics"
            ],
            "Mathematics": [
                "Number Theory", "Algebra", "Geometry", "Analysis",
                "Topology", "Applied Mathematics", "Statistics",
                "Differential Equations", "Graph Theory", "Cryptography",
                "Mathematical Logic", "Probability Theory", "Numerical Analysis",
                "Optimization", "Game Theory", "Category Theory",
                "Combinatorics", "Complex Analysis", "Dynamical Systems",
                "Algebraic Geometry", "Functional Analysis", "Mathematical Physics"
            ],
            "Chemistry": [
                "Organic Chemistry", "Inorganic Chemistry", "Physical Chemistry",
                "Analytical Chemistry", "Biochemistry", "Materials Science",
                "Polymer Chemistry", "Medicinal Chemistry", "Computational Chemistry",
                "Environmental Chemistry", "Nuclear Chemistry", "Electrochemistry",
                "Green Chemistry", "Nanotechnology", "Photochemistry",
                "Quantum Chemistry", "Surface Chemistry", "Catalysis",
                "Spectroscopy", "Chemical Engineering", "Crystallography"
            ],
            "Biology": [
                "Molecular Biology", "Genetics", "Ecology", "Neuroscience",
                "Microbiology", "Biotechnology", "Cell Biology", "Immunology",
                "Evolutionary Biology", "Marine Biology", "Plant Biology",
                "Zoology", "Biochemistry", "Bioinformatics", "Systems Biology",
                "Cancer Biology", "Developmental Biology", "Structural Biology",
                "Genomics", "Proteomics", "Stem Cell Research",
                "Environmental Biology", "Conservation Biology", "Virology"
            ],
            "Engineering": [
                "Mechanical Engineering", "Electrical Engineering",
                "Civil Engineering", "Chemical Engineering", "Robotics",
                "Aerospace Engineering", "Biomedical Engineering",
                "Environmental Engineering", "Materials Engineering",
                "Industrial Engineering", "Nuclear Engineering",
                "Automotive Engineering", "Marine Engineering",
                "Mechatronics", "Control Systems", "Telecommunications",
                "Structural Engineering", "Power Systems", "MEMS",
                "Nanotechnology Engineering", "Software Engineering",
                "Systems Engineering", "Manufacturing Engineering"
            ],
            "Environmental Science": [
                "Climate Change", "Sustainability", "Environmental Protection",
                "Renewable Energy", "Ecology", "Conservation",
                "Environmental Chemistry", "Atmospheric Science",
                "Marine Science", "Soil Science", "Waste Management",
                "Environmental Policy", "Biodiversity", "Environmental Impact Assessment",
                "Natural Resource Management", "Environmental Toxicology",
                "Water Resources", "Environmental Biotechnology",
                "Sustainable Agriculture", "Green Technology"
            ],
            "Medicine": [
                "Public Health", "Pharmacology", "Immunology",
                "Medical Technology", "Clinical Research", "Cardiology",
                "Neurology", "Oncology", "Pediatrics", "Surgery",
                "Psychiatry", "Radiology", "Emergency Medicine",
                "Endocrinology", "Genetics", "Pathology", "Dermatology",
                "Ophthalmology", "Orthopedics", "Internal Medicine",
                "Infectious Diseases", "Regenerative Medicine"
            ],
            "Psychology": [
                "Clinical Psychology", "Cognitive Psychology", "Developmental Psychology",
                "Social Psychology", "Behavioral Psychology", "Neuropsychology",
                "Educational Psychology", "Industrial Psychology", "Health Psychology",
                "Forensic Psychology", "Sports Psychology", "Positive Psychology",
                "Experimental Psychology", "Personality Psychology", "Psychotherapy",
                "Child Psychology", "Abnormal Psychology", "Environmental Psychology"
            ],
            "Economics": [
                "Microeconomics", "Macroeconomics", "Behavioral Economics",
                "Financial Economics", "International Economics", "Development Economics",
                "Environmental Economics", "Labor Economics", "Industrial Organization",
                "Public Economics", "Health Economics", "Agricultural Economics",
                "Monetary Economics", "Game Theory", "Econometrics",
                "Urban Economics", "Political Economy", "Digital Economics"
            ],
            "Business": [
                "Strategic Management", "Marketing", "Finance",
                "Operations Management", "Human Resources", "International Business",
                "Entrepreneurship", "Supply Chain Management", "Business Analytics",
                "Digital Marketing", "Risk Management", "Project Management",
                "Corporate Governance", "Business Ethics", "Innovation Management",
                "Knowledge Management", "Change Management", "Quality Management"
            ],
            "Social Sciences": [
                "Sociology", "Anthropology", "Political Science",
                "International Relations", "Public Policy", "Urban Planning",
                "Development Studies", "Gender Studies", "Cultural Studies",
                "Human Geography", "Demography", "Social Work",
                "Communication Studies", "Criminology", "Education",
                "Library Science", "Media Studies", "Social Psychology"
            ],
            "Arts and Humanities": [
                "Literature", "Philosophy", "History",
                "Art History", "Linguistics", "Religious Studies",
                "Music Theory", "Film Studies", "Theater Studies",
                "Classical Studies", "Digital Humanities", "Archaeology",
                "Cultural Heritage", "Ethics", "Visual Arts",
                "Comparative Literature", "Translation Studies", "Performance Studies"
            ],
            "Information Science": [
                "Data Science", "Information Systems", "Digital Libraries",
                "Knowledge Management", "Information Security", "Database Management",
                "Information Architecture", "Digital Preservation", "Information Retrieval",
                "Metadata Management", "Information Policy", "Social Informatics",
                "Web Science", "Information Ethics", "Digital Curation"
            ]
        }

        # 随机选择一个领域和子领域
        field = random.choice(list(fields.keys()))
        subfield = random.choice(fields[field])

        prompt = f"""Generate ONE unique research topic in the field of {field}, specifically focusing on {subfield}.
        Requirements:
        1. The topic should be specific and well-defined (not too broad, not too narrow)
        2. It should be suitable for academic research
        3. It should reflect current trends and challenges in {subfield}
        4. Identify 2-3 key aspects that the research should focus on
        
        Respond with a JSON object in this format:
        {{
            "field": "{field}",
            "subfield": "{subfield}",
            "topic": "Your specific research topic here",
            "key_aspects": ["aspect1", "aspect2"]
        }}"""

        messages = [
            {
                "role": "system",
                "content": "You are a research topic generator. Generate a meaningful academic research topic. Respond with a JSON object"
            },
            {"role": "user", "content": prompt}
        ]

        try:
            response = self._call_api(messages, temperature=0.5,response_format="json_object")
            topic = json.loads(response.strip())
            # 验证返回的field和subfield是否匹配我们的随机选择
            if topic["field"] != field or topic["subfield"] != subfield:
                topic["field"] = field
                topic["subfield"] = subfield
            
            return topic
        except Exception as e:
            print(f"生成主题时出错: {str(e)}")
            exit(1)

    def batch_generate_papers(self, count: int = 100, output_dir: str = "generated_papers"):
        """批量生成多篇论文"""
        # 创建输出目录
        os.makedirs(output_dir, exist_ok=True)
        
        # 创建进度记录文件
        progress_file = os.path.join(output_dir, "generation_progress.json")
        if os.path.exists(progress_file):
            with open(progress_file, 'r', encoding='utf-8') as f:
                progress = json.load(f)
        else:
            progress = {"completed": [], "failed": [], "count": 0}
        
        # 继续生成，直到达到目标数量
        remaining = count - len(progress["completed"])
        if remaining <= 0:
            print(f"已完成指定数量的论文生成({count}篇)")
            return
        
        print(f"继续生成剩余的{remaining}篇论文...")
        
        # 逐个生成论文
        for i in range(remaining):
            # 生成新主题
            topic_info = self.generate_single_topic()
            topic = topic_info["topic"]
            field = topic_info["field"]
            
            # 检查是否已经生成
            if topic in progress["completed"]:
                print(f"跳过重复的主题: {topic}")
                continue
            
            try:
                current_count = len(progress["completed"]) + 1
                print(f"\n[{current_count}/{count}] 正在生成论文: {topic} (领域: {field})")
                
                # 生成论文
                self.topic = topic
                
                # 创建子目录
                field_dir = os.path.join(output_dir, field.replace(" ", "_"))
                os.makedirs(field_dir, exist_ok=True)
                
                # 生成文件名
                model_name = self.current_model  # 获取当前使用的模型名称
                paper_count = len(progress["completed"]) + 1  # 获取当前是第几篇论文
                filename = f"{topic[:10].replace(' ', '_')}-{model_name}-{paper_count}.pdf"
                filepath = os.path.join(field_dir, filename)
                
                # 生成论文并保存为PDF
                self.save_paper_to_pdf(filepath)
                
                # 更新进度
                progress["completed"].append(topic)
                progress["count"] = len(progress["completed"])
                
                # 保存进度
                with open(progress_file, 'w', encoding='utf-8') as f:
                    json.dump(progress, f, indent=2)
                    
                print(f"论文已保存: {filepath}")
                
                # 添加随机延迟避免API限制
                time.sleep(random.uniform(5, 10))
                
            except Exception as e:
                print(f"生成论文失败: {topic}")
                print(f"错误信息: {str(e)}")
                progress["failed"].append({"topic": topic, "error": str(e)})
                
                # 保存进度
                with open(progress_file, 'w', encoding='utf-8') as f:
                    json.dump(progress, f, indent=2)
                
                # 继续下一篇
                continue

        print("\n批量生成完成!")
        print(f"成功: {len(progress['completed'])} 篇")
        print(f"失败: {len(progress['failed'])} 篇")
        
        # 保存失败记录
        if progress["failed"]:
            failed_file = os.path.join(output_dir, "failed_papers.json")
            with open(failed_file, 'w', encoding='utf-8') as f:
                json.dump(progress["failed"], f, indent=2)
            print(f"失败记录已保存到: {failed_file}")

if __name__ == "__main__":
    # 创建生成器实例，默认使用gpt4o
    generator = AdvancedPaperGenerator(model_name='doubao')
    
    # 示例1：使用单个模型生成论文
    #generator.save_paper_to_pdf("generated_paper1.pdf")
    
    # 示例2：使用多个模型生成论文
    generator.batch_generate_papers(
        count=900,  
        output_dir="batch_generated_papers",
    )
    
    # 示例3：手动切换模型
    #generator.switch_model('qwen')
    #generator.batch_generate_papers(count=3, output_dir="qwen_papers")






















